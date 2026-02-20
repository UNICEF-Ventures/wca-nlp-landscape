"""Generate DOCX documents from research YAML data.

Usage:
    python scripts/generate_docs.py              # Both documents
    python scripts/generate_docs.py --languages   # Languages only
    python scripts/generate_docs.py --actors      # Actors only
"""

import argparse
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# Add scripts/ to path so we can import htmlgen
sys.path.insert(0, str(Path(__file__).parent))

from htmlgen.data import load_all_languages, load_all_actors, load_focused_languages
from htmlgen.utils import (
    format_actor_type, country_codes_to_names, get_model_count,
    format_number, get_actors_for_language, TASK_LABELS, TASK_SHORT_LABELS,
)
from htmlgen.constants import WCA_COUNTRIES, OUTPUT_DIR, COUNTRY_NAMES

REPO_URL = "https://github.com/translatorswb/wca-nlp-landscape"
LIVE_URL = "https://translatorswb.github.io/wca-nlp-landscape/"


# --- Styling helpers ---

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)


def style_header_row(row, bg_color='1B3A5C'):
    """Style a table header row with background color and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(8)


def set_table_style(table):
    """Apply compact styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    if not run.font.size:
                        run.font.size = Pt(8)


def add_header(doc, title, generated_line, source_url, live_url):
    """Add document header with title, generated date, and clickable links."""
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Generated timestamp
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(generated_line)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Source repo link
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Source: ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_hyperlink(p, source_url, source_url, font_size=Pt(9))

    # Live site link
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Live site: ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_hyperlink(p, live_url, live_url, font_size=Pt(9))


def add_hyperlink(paragraph, text, url, font_size=Pt(8), bold=False):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('r:id'): r_id})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    # Blue underlined style
    color_elem = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0066CC'})
    rPr.append(color_elem)
    u_elem = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(u_elem)
    sz_val = str(int(font_size.pt * 2))  # half-points
    sz = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): sz_val})
    rPr.append(sz)
    szCs = paragraph._element.makeelement(qn('w:szCs'), {qn('w:val'): sz_val})
    rPr.append(szCs)
    if bold:
        b_elem = paragraph._element.makeelement(qn('w:b'), {})
        rPr.append(b_elem)
    new_run.append(rPr)
    text_elem = paragraph._element.makeelement(qn('w:t'), {})
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_bookmark(paragraph, bookmark_name):
    """Add an invisible bookmark anchor at the current paragraph."""
    # Use a unique integer id based on the bookmark name hash
    bm_id = str(abs(hash(bookmark_name)) % 1000000)
    start = paragraph._element.makeelement(qn('w:bookmarkStart'), {
        qn('w:id'): bm_id,
        qn('w:name'): bookmark_name,
    })
    end = paragraph._element.makeelement(qn('w:bookmarkEnd'), {
        qn('w:id'): bm_id,
    })
    paragraph._element.append(start)
    paragraph._element.append(end)


def add_internal_link(paragraph, text, bookmark_name, font_size=Pt(8), bold=False):
    """Add a clickable internal link (to a bookmark within the same document)."""
    hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {qn('w:anchor'): bookmark_name})
    new_run = paragraph._element.makeelement(qn('w:r'), {})
    rPr = paragraph._element.makeelement(qn('w:rPr'), {})
    color_elem = paragraph._element.makeelement(qn('w:color'), {qn('w:val'): '0066CC'})
    rPr.append(color_elem)
    u_elem = paragraph._element.makeelement(qn('w:u'), {qn('w:val'): 'single'})
    rPr.append(u_elem)
    sz_val = str(int(font_size.pt * 2))
    sz = paragraph._element.makeelement(qn('w:sz'), {qn('w:val'): sz_val})
    rPr.append(sz)
    szCs = paragraph._element.makeelement(qn('w:szCs'), {qn('w:val'): sz_val})
    rPr.append(szCs)
    if bold:
        b_elem = paragraph._element.makeelement(qn('w:b'), {})
        rPr.append(b_elem)
    new_run.append(rPr)
    text_elem = paragraph._element.makeelement(qn('w:t'), {})
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def set_cell_link(cell, text, url, font_size=Pt(8)):
    """Set a table cell's content to a clickable hyperlink."""
    cell.text = ''
    add_hyperlink(cell.paragraphs[0], text, url, font_size=font_size)


def set_cell_internal_link(cell, text, bookmark_name, font_size=Pt(8)):
    """Set a table cell's content to a clickable internal link."""
    cell.text = ''
    add_internal_link(cell.paragraphs[0], text, bookmark_name, font_size=font_size)


def get_wca_countries(countries_list):
    """Filter country list to only WCA countries."""
    return [c for c in countries_list if c in WCA_COUNTRIES]


def get_total_hf_models(models_data):
    """Sum model counts across all tasks."""
    total = 0
    for task in ['asr', 'tts', 'translation', 'llm']:
        total += get_model_count(models_data.get(task, {}))
    return total


def get_hf_search_url(iso1, iso3, pipeline_tag):
    """Build HuggingFace model search URL."""
    code = iso1 if iso1 else iso3
    if not code:
        return None
    return f"https://huggingface.co/models?pipeline_tag={pipeline_tag}&language={code}&sort=trending"


def get_hf_datasets_url(iso1, iso3, task_category):
    """Build HuggingFace dataset search URL."""
    code = iso1 if iso1 else iso3
    if not code:
        return None
    return f"https://huggingface.co/datasets?task_categories=task_categories:{task_category}&language=language:{code}&sort=trending"


# --- Languages document ---

def generate_languages_doc(languages, actors):
    """Generate the Languages DOCX document."""
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    now = datetime.now(timezone.utc)
    add_header(doc, "WCA NLP Landscape \u2014 Languages",
               f"Generated: {now.strftime('%Y-%m-%d %H:%M')} UTC",
               REPO_URL, LIVE_URL)

    # Sort languages alphabetically by name
    sorted_langs = sorted(
        languages.items(),
        key=lambda x: x[1].get('info', {}).get('name', x[0]).lower()
    )

    # --- Summary Matrix ---
    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph(f"{len(sorted_langs)} focus languages covered.")
    p.runs[0].font.size = Pt(9)

    cols = ['Language', 'ISO', 'Countries', 'Family', 'Speakers (L1/L2)', 'HF Models', 'Benchmarks']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'

    # Header
    for i, col_name in enumerate(cols):
        table.rows[0].cells[i].text = col_name
    style_header_row(table.rows[0])

    for iso_code, lang_data in sorted_langs:
        info = lang_data.get('info', {})
        models = lang_data.get('models', {})
        benchmarks = lang_data.get('benchmarks', {})

        name = info.get('name', iso_code)
        wiki = info.get('wikipedia', {}) or {}
        family = wiki.get('family', '')
        speakers_l1 = wiki.get('speakers_l1', '')
        speakers_l2 = wiki.get('speakers_l2', '')

        # Format speakers
        if speakers_l1 and speakers_l2:
            speakers = f"{speakers_l1} / {speakers_l2}"
        elif speakers_l1:
            speakers = str(speakers_l1)
        else:
            speakers = '—'

        countries = info.get('countries', [])
        wca = get_wca_countries(countries)
        countries_str = ', '.join(wca[:5])
        if len(wca) > 5:
            countries_str += f' +{len(wca)-5}'

        total_models = get_total_hf_models(models)
        has_benchmarks = bool(benchmarks.get('evaluations'))

        row = table.add_row()
        # Language name as internal link to its section
        bookmark_name = f"lang_{iso_code}"
        set_cell_internal_link(row.cells[0], name, bookmark_name)
        # Rest as plain text
        row.cells[1].text = iso_code
        row.cells[2].text = countries_str
        row.cells[3].text = family or '—'
        row.cells[4].text = speakers
        row.cells[5].text = str(total_models)
        row.cells[6].text = '\u2713' if has_benchmarks else '\u2717'

    set_table_style(table)
    doc.add_page_break()

    # --- Language Entries ---
    for lang_idx, (iso_code, lang_data) in enumerate(sorted_langs):
        info = lang_data.get('info', {})
        models = lang_data.get('models', {})
        datasets = lang_data.get('datasets', {})
        benchmarks = lang_data.get('benchmarks', {})

        name = info.get('name', iso_code)
        iso3 = info.get('iso_639_3', iso_code)
        iso1 = info.get('iso_639_1', '')
        name_french = info.get('name_french', '')
        altnames = info.get('altnames', [])
        countries = info.get('countries', [])
        population = info.get('population', '')
        endangerment = info.get('endangerment', '')

        wiki = info.get('wikipedia', {}) or {}
        wiki_family = wiki.get('family', '')
        wiki_speakers_l1 = wiki.get('speakers_l1', '')
        wiki_speakers_l2 = wiki.get('speakers_l2', '')
        wiki_writing = wiki.get('writing_system', '')
        wiki_url = wiki.get('url', '')

        tech_resources = info.get('tech_resources', {}) or {}
        resource_links = info.get('resource_links', {}) or {}

        # --- Language heading with bookmark ---
        heading = doc.add_heading(f"{name} ({iso_code})", level=1)
        add_bookmark(heading, f"lang_{iso_code}")

        # --- General Information ---
        doc.add_heading("General Information", level=2)

        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'

        # Plain text info items
        info_items = [
            ('ISO 639-3', iso3),
            ('ISO 639-1', iso1 or '—'),
            ('French Name', name_french or '—'),
            ('Language Family', wiki_family or '—'),
            ('Writing System', wiki_writing or '—'),
            ('L1 Speakers', wiki_speakers_l1 or '—'),
            ('L2 Speakers', wiki_speakers_l2 or '—'),
            ('Population', f"{population}" if population else '—'),
            ('Endangerment', endangerment or '—'),
            ('Countries', ', '.join(countries) if countries else '—'),
        ]
        if altnames:
            info_items.append(('Also Known As', ', '.join(altnames[:10])))

        for label, value in info_items:
            row = info_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(8)

        # Wikipedia as clickable link
        if wiki_url:
            row = info_table.add_row()
            row.cells[0].text = 'Wikipedia'
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            set_cell_link(row.cells[1], f"{name} on Wikipedia", wiki_url)

        set_table_style(info_table)

        # --- Tech Resources ---
        total_models = get_total_hf_models(models)
        if tech_resources or resource_links or total_models > 0:
            doc.add_heading("NLP & Tech Resources", level=2)

            if total_models > 0:
                p = doc.add_paragraph(f"HuggingFace Models: {total_models} total", style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(8)

            for k, url in resource_links.items():
                if url:
                    p = doc.add_paragraph(style='List Bullet')
                    add_hyperlink(p, k, url)

            for k, v in tech_resources.items():
                if v is True:
                    p = doc.add_paragraph(f"\u2713 {k}", style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(8)
                elif isinstance(v, str) and v.startswith('http'):
                    p = doc.add_paragraph(style='List Bullet')
                    add_hyperlink(p, k, v)
                elif v:
                    p = doc.add_paragraph(f"\u2713 {k}: {v}", style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(8)

        # --- Actors ---
        lang_actors = get_actors_for_language(iso_code, actors)
        if lang_actors:
            doc.add_heading(f"Actors ({len(lang_actors)})", level=2)
            for actor in lang_actors:
                actor_name = actor.get('name', '')
                actor_website = actor.get('website', '')
                actor_type = format_actor_type(actor.get('type', ''))
                p = doc.add_paragraph()
                if actor_website:
                    add_hyperlink(p, actor_name, actor_website, font_size=Pt(9), bold=True)
                else:
                    run = p.add_run(actor_name)
                    run.bold = True
                    run.font.size = Pt(9)
                run = p.add_run(f" ({actor_type})")
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # --- Common Voice ---
        cv = benchmarks.get('common_voice') or {}
        if cv:
            doc.add_heading("Common Voice", level=2)
            cv_table = doc.add_table(rows=0, cols=2)
            cv_table.style = 'Table Grid'
            cv_items = [
                ('Total Hours', str(cv.get('total_hours', 0))),
                ('Total Clips', str(cv.get('total_clips', 0))),
                ('Validated Clips', str(cv.get('validated_clips', 0))),
                ('Gender', f"Male {cv.get('male_percent', 0)}% / Female {cv.get('female_percent', 0)}%"),
            ]
            for label, value in cv_items:
                row = cv_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = value
                for run in row.cells[0].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(8)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(8)
            set_table_style(cv_table)

        # --- Benchmarks ---
        evaluations = benchmarks.get('evaluations', {})
        if evaluations:
            doc.add_heading("Benchmarks", level=2)
            _render_benchmarks_docx(doc, evaluations)

        # --- Unbenchmarked Models ---
        unbenchmarked = benchmarks.get('unbenchmarked_models', [])
        if unbenchmarked:
            doc.add_heading("Noteworthy Models Without Benchmark", level=2)
            ub_table = doc.add_table(rows=1, cols=3)
            ub_table.style = 'Table Grid'
            ub_table.rows[0].cells[0].text = 'Model'
            ub_table.rows[0].cells[1].text = 'Task'
            ub_table.rows[0].cells[2].text = 'Notes'
            style_header_row(ub_table.rows[0])

            for entry in unbenchmarked:
                model = entry.get('model', '')
                model_url = entry.get('model_url', '')
                task = entry.get('task', '')
                notes = entry.get('notes', '')
                task_label = TASK_SHORT_LABELS.get(task, task.upper())
                row = ub_table.add_row()
                if model_url:
                    set_cell_link(row.cells[0], model, model_url)
                else:
                    row.cells[0].text = model
                row.cells[1].text = task_label
                row.cells[2].text = notes
            set_table_style(ub_table)

        # --- Models & Datasets Summary ---
        doc.add_heading("Models & Datasets on HuggingFace", level=2)

        pipeline_tags = {
            'asr': 'automatic-speech-recognition',
            'tts': 'text-to-speech',
            'translation': 'translation',
            'llm': 'text-generation',
        }
        dataset_categories = {
            'asr': 'automatic-speech-recognition',
            'tts': 'text-to-speech',
        }

        for task_key, pipeline_tag in pipeline_tags.items():
            task_label = TASK_LABELS.get(task_key, task_key.upper())
            task_data = models.get(task_key, {})
            total_count = get_model_count(task_data)

            p = doc.add_paragraph()
            run = p.add_run(f"{task_label}: ")
            run.bold = True
            run.font.size = Pt(8)

            # Check if we have counts_by_code data (new format)
            if isinstance(task_data, dict) and 'counts_by_code' in task_data:
                counts_by_code = task_data['counts_by_code']
                if counts_by_code:
                    # Generate links for each code (ISO-2 and ISO-3)
                    first = True
                    for code, count in sorted(counts_by_code.items(), key=lambda x: -x[1]):
                        if count > 0:
                            if not first:
                                run = p.add_run(" | ")
                                run.font.size = Pt(8)
                            url = f"https://huggingface.co/models?pipeline_tag={pipeline_tag}&language={code}&sort=trending"
                            add_hyperlink(p, f"{code}: {count}", url)
                            first = False
                else:
                    run = p.add_run(f"{total_count} models")
                    run.font.size = Pt(8)
            else:
                # Fallback to single link
                url = get_hf_search_url(iso1, iso3, pipeline_tag)
                if url and total_count > 0:
                    add_hyperlink(p, f"{total_count} models", url)
                else:
                    run = p.add_run(f"{total_count} models")
                    run.font.size = Pt(8)

        for task_key, task_cat in dataset_categories.items():
            task_label = TASK_LABELS.get(task_key, task_key.upper())
            task_data = datasets.get(task_key, {})
            total_count = get_model_count(task_data)

            p = doc.add_paragraph()
            run = p.add_run(f"{task_label} Datasets: ")
            run.bold = True
            run.font.size = Pt(8)

            # Check if we have counts_by_code data (new format)
            if isinstance(task_data, dict) and 'counts_by_code' in task_data:
                counts_by_code = task_data['counts_by_code']
                if counts_by_code:
                    # Generate links for each code (ISO-2 and ISO-3)
                    first = True
                    for code, count in sorted(counts_by_code.items(), key=lambda x: -x[1]):
                        if count > 0:
                            if not first:
                                run = p.add_run(" | ")
                                run.font.size = Pt(8)
                            url = f"https://huggingface.co/datasets?task_categories=task_categories:{task_cat}&language=language:{code}&sort=trending"
                            add_hyperlink(p, f"{code}: {count}", url)
                            first = False
                else:
                    run = p.add_run(f"{total_count} datasets")
                    run.font.size = Pt(8)
            else:
                # Fallback to single link
                url = get_hf_datasets_url(iso1, iso3, task_cat)
                if url and total_count > 0:
                    add_hyperlink(p, f"{total_count} datasets", url)
                else:
                    run = p.add_run(f"{total_count} datasets")
                    run.font.size = Pt(8)

        # Page break between languages (except last)
        if lang_idx < len(sorted_langs) - 1:
            doc.add_page_break()

    return doc


def _render_benchmarks_docx(doc, evaluations):
    """Render benchmark tables for all tasks into the document."""
    task_order = ['asr', 'tts', 'mt', 'llm']

    for task in task_order:
        entries = evaluations.get(task, [])
        if entries:
            label = TASK_LABELS.get(task, task.upper())
            _render_single_benchmark_table(doc, label, entries)

    # Non-standard tasks
    for task, entries in evaluations.items():
        if task in TASK_LABELS or not entries:
            continue
        _render_single_benchmark_table(doc, task.upper(), entries)


def _render_single_benchmark_table(doc, label, entries):
    """Render a single benchmark task table."""
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(9)

    # Collect all unique metric names
    metric_names = []
    seen = set()
    for entry in entries:
        for result in entry.get('results', []):
            for m in result.get('metrics', []):
                mname = m.get('name', '')
                if mname and mname not in seen:
                    metric_names.append(mname)
                    seen.add(mname)

    # Build table
    col_headers = ['Model', 'Test Set'] + metric_names + ['Source']
    table = doc.add_table(rows=1, cols=len(col_headers))
    table.style = 'Table Grid'

    for i, header in enumerate(col_headers):
        table.rows[0].cells[i].text = header
    style_header_row(table.rows[0])

    for entry in entries:
        model = entry.get('model', '')
        model_url = entry.get('model_url', '')
        results = entry.get('results', [])

        if not results:
            row = table.add_row()
            if model_url:
                set_cell_link(row.cells[0], model, model_url)
            else:
                row.cells[0].text = model
            row.cells[1].text = '—'
            for i in range(len(metric_names)):
                row.cells[2 + i].text = '—'
            row.cells[-1].text = '—'
            continue

        for result in results:
            test_set = result.get('test_set', '')
            source = result.get('source', '')
            source_url = result.get('source_url', '')

            metrics_by_name = {
                m['name']: m['value']
                for m in result.get('metrics', [])
                if 'name' in m
            }

            row = table.add_row()
            # Model name as clickable HuggingFace link
            if model_url:
                set_cell_link(row.cells[0], model, model_url)
            else:
                row.cells[0].text = model
            row.cells[1].text = test_set
            for i, mname in enumerate(metric_names):
                val = metrics_by_name.get(mname)
                row.cells[2 + i].text = str(val) if val is not None else '—'

            # Source as clickable link
            if source_url:
                set_cell_link(row.cells[-1], source or 'link', source_url)
            else:
                row.cells[-1].text = source or '—'

    set_table_style(table)


# --- Actors document ---

def generate_actors_doc(actors, languages):
    """Generate the Actors DOCX document."""
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    now = datetime.now(timezone.utc)
    add_header(doc, "WCA NLP Landscape \u2014 Actors",
               f"Generated: {now.strftime('%Y-%m-%d %H:%M')} UTC",
               REPO_URL, LIVE_URL)

    # Filter out template
    actor_items = [
        (k, v) for k, v in sorted(actors.items(), key=lambda x: x[1].get('name', x[0]).lower())
        if k != 'actor-template'
    ]

    # --- Summary Matrix ---
    doc.add_heading("Summary", level=1)
    p = doc.add_paragraph(f"{len(actor_items)} actors mapped.")
    p.runs[0].font.size = Pt(9)

    cols = ['Actor', 'Type', 'Countries', 'Languages', 'Openness', 'Maturity', 'Engagement']
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'

    for i, col_name in enumerate(cols):
        table.rows[0].cells[i].text = col_name
    style_header_row(table.rows[0])

    for actor_key, actor_data in actor_items:
        name = actor_data.get('name', actor_key)
        actor_type = format_actor_type(actor_data.get('type', ''))
        actor_countries = actor_data.get('countries', [])
        actor_langs = actor_data.get('languages', [])
        openness = actor_data.get('openness', '—') or '—'
        maturity = actor_data.get('maturity', '—') or '—'
        engagement = actor_data.get('engagement_status', '—') or '—'

        country_names = country_codes_to_names(actor_countries)
        countries_str = ', '.join(country_names[:5])
        if len(country_names) > 5:
            countries_str += f' +{len(country_names)-5}'

        langs_str = ', '.join(str(l) for l in actor_langs[:8])
        if len(actor_langs) > 8:
            langs_str += f' +{len(actor_langs)-8}'

        row = table.add_row()
        # Actor name as internal link to its section (+ website fallback)
        website = actor_data.get('website', '')
        bookmark_name = f"actor_{actor_data.get('id', actor_key)}"
        set_cell_internal_link(row.cells[0], name, bookmark_name)
        row.cells[1].text = actor_type
        row.cells[2].text = countries_str
        row.cells[3].text = langs_str
        row.cells[4].text = openness
        row.cells[5].text = maturity
        row.cells[6].text = engagement

    set_table_style(table)
    doc.add_page_break()

    # Build ISO-to-name mapping for language references
    iso_to_name = {}
    for iso, ldata in languages.items():
        linfo = ldata.get('info', {})
        iso_to_name[iso] = linfo.get('name', iso)

    # --- Actor Entries ---
    for actor_idx, (actor_key, actor_data) in enumerate(actor_items):
        name = actor_data.get('name', actor_key)
        actor_type = format_actor_type(actor_data.get('type', ''))
        website = actor_data.get('website', '')
        location = actor_data.get('location', '')
        founded = actor_data.get('founded', '')
        organization_size = actor_data.get('organization_size', '')
        funding = actor_data.get('funding', '')
        maturity = actor_data.get('maturity', '')
        openness = actor_data.get('openness', '')
        engagement_status = actor_data.get('engagement_status', '')
        description = actor_data.get('description', '')
        unicef_relevance = actor_data.get('unicef_relevance', '')
        github = actor_data.get('github', '')
        huggingface = actor_data.get('huggingface', '')
        actor_countries = actor_data.get('countries', [])
        actor_langs = actor_data.get('languages', [])
        key_people = actor_data.get('key_people', [])
        projects = actor_data.get('projects', [])
        publications = actor_data.get('publications', [])
        partnerships = actor_data.get('partnerships', [])
        notes = actor_data.get('notes', '')
        last_updated = actor_data.get('last_updated', '')

        # Contact
        contact = actor_data.get('contact', '')
        if isinstance(contact, dict):
            contact_parts = []
            if contact.get('email'):
                contact_parts.append(contact['email'])
            if contact.get('mailing_list'):
                contact_parts.append(f"Mailing list: {contact['mailing_list']}")
            if contact.get('slack'):
                contact_parts.append(f"Slack: {contact['slack']}")
            contact_str = ' | '.join(contact_parts) if contact_parts else '—'
        else:
            contact_str = str(contact) if contact else '—'

        # --- Actor heading with bookmark ---
        actor_id = actor_data.get('id', actor_key)
        heading = doc.add_heading(name, level=1)
        add_bookmark(heading, f"actor_{actor_id}")

        # --- Info table ---
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'

        # Plain text info items
        plain_items = [
            ('Type', actor_type),
            ('Founded', str(founded) if founded else '—'),
            ('Size', organization_size or '—'),
            ('Maturity', maturity or '—'),
            ('Openness', openness or '—'),
            ('Engagement', engagement_status or '—'),
            ('Location', location or '—'),
            ('Funding', funding or '—'),
        ]

        for label, value in plain_items:
            row = info_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(8)

        # Link items: Website, GitHub, HuggingFace
        link_items = [
            ('Website', website),
            ('GitHub', github),
            ('HuggingFace', huggingface),
        ]
        for label, url in link_items:
            row = info_table.add_row()
            row.cells[0].text = label
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)
            if url:
                set_cell_link(row.cells[1], url, url)
            else:
                row.cells[1].text = '—'
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(8)

        # Contact row
        row = info_table.add_row()
        row.cells[0].text = 'Contact'
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        # Contact may have emails/links mixed
        contact_raw = actor_data.get('contact', '')
        if isinstance(contact_raw, dict):
            cp = row.cells[1].paragraphs[0]
            parts_added = False
            if contact_raw.get('email'):
                email = contact_raw['email']
                add_hyperlink(cp, email, f"mailto:{email}")
                parts_added = True
            if contact_raw.get('mailing_list'):
                if parts_added:
                    run = cp.add_run(' | ')
                    run.font.size = Pt(8)
                add_hyperlink(cp, 'Mailing list', contact_raw['mailing_list'])
                parts_added = True
            if contact_raw.get('slack'):
                if parts_added:
                    run = cp.add_run(' | ')
                    run.font.size = Pt(8)
                add_hyperlink(cp, 'Slack', contact_raw['slack'])
                parts_added = True
            if not parts_added:
                row.cells[1].text = '—'
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(8)
        elif contact_raw and '@' in str(contact_raw):
            set_cell_link(row.cells[1], str(contact_raw), f"mailto:{contact_raw}")
        else:
            row.cells[1].text = str(contact_raw) if contact_raw else '—'
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(8)

        set_table_style(info_table)

        # --- Description ---
        if description:
            doc.add_heading("Description", level=2)
            p = doc.add_paragraph(str(description).strip())
            for run in p.runs:
                run.font.size = Pt(9)

        # --- UNICEF Relevance ---
        if unicef_relevance:
            doc.add_heading("UNICEF Relevance", level=2)
            p = doc.add_paragraph(str(unicef_relevance).strip())
            for run in p.runs:
                run.font.size = Pt(9)

        # --- Countries & Languages ---
        country_names = country_codes_to_names(actor_countries)
        if country_names or actor_langs:
            doc.add_heading("Coverage", level=2)
            if country_names:
                p = doc.add_paragraph()
                run = p.add_run("Countries: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(', '.join(country_names))
                run.font.size = Pt(9)
            if actor_langs:
                lang_display = []
                for lcode in actor_langs:
                    lname = iso_to_name.get(str(lcode))
                    if lname:
                        lang_display.append(f"{lname} ({lcode})")
                    else:
                        lang_display.append(str(lcode))
                p = doc.add_paragraph()
                run = p.add_run("Languages: ")
                run.bold = True
                run.font.size = Pt(9)
                run = p.add_run(', '.join(lang_display))
                run.font.size = Pt(9)

        # --- Key People ---
        if key_people:
            doc.add_heading("Key People", level=2)
            for person in key_people:
                if isinstance(person, dict):
                    pname = person.get('name', '')
                    prole = person.get('role', '')
                    paffil = person.get('affiliation', '')
                    pnote = person.get('note', '')
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(pname)
                    run.bold = True
                    run.font.size = Pt(9)
                    role_str = f" — {prole}"
                    if paffil:
                        role_str += f" ({paffil})"
                    run = p.add_run(role_str)
                    run.font.size = Pt(8)
                    if pnote:
                        run = p.add_run(f"\n{pnote}")
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    p = doc.add_paragraph(str(person), style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(9)

        # --- Projects ---
        if projects:
            projects_url = actor_data.get('projects_url', '')
            doc.add_heading("Projects", level=2)
            for proj in projects:
                if isinstance(proj, dict):
                    pname = proj.get('name', '')
                    pdesc = proj.get('description', '')
                    purl = proj.get('url', '')
                    plangs = proj.get('languages', [])
                    p = doc.add_paragraph(style='List Bullet')
                    if purl:
                        add_hyperlink(p, pname, purl, font_size=Pt(9), bold=True)
                    else:
                        run = p.add_run(pname)
                        run.bold = True
                        run.font.size = Pt(9)
                    if pdesc:
                        run = p.add_run(f" — {pdesc}")
                        run.font.size = Pt(8)
                    if plangs:
                        lang_names = [iso_to_name.get(str(l), str(l)) for l in plangs]
                        run = p.add_run(f"\nLanguages: {', '.join(lang_names)}")
                        run.font.size = Pt(7)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    p = doc.add_paragraph(str(proj), style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(9)
            if projects_url:
                p = doc.add_paragraph()
                add_hyperlink(p, "View all projects \u2192", projects_url, font_size=Pt(8))

        # --- Publications ---
        if publications:
            publications_url = actor_data.get('publications_url', '')
            doc.add_heading("Publications", level=2)
            for pub in publications:
                if isinstance(pub, dict):
                    ptitle = pub.get('title', '')
                    pvenue = pub.get('venue', '')
                    purl = pub.get('url', '')
                    p = doc.add_paragraph(style='List Bullet')
                    if purl:
                        add_hyperlink(p, ptitle, purl, font_size=Pt(9))
                    else:
                        run = p.add_run(ptitle)
                        run.font.size = Pt(9)
                    if pvenue:
                        run = p.add_run(f" ({pvenue})")
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    p = doc.add_paragraph(str(pub), style='List Bullet')
                    for run in p.runs:
                        run.font.size = Pt(9)
            if publications_url:
                p = doc.add_paragraph()
                add_hyperlink(p, "View all publications \u2192", publications_url, font_size=Pt(8))

        # --- Partnerships ---
        if partnerships:
            doc.add_heading("Partnerships", level=2)
            for partner in partnerships:
                p = doc.add_paragraph(str(partner), style='List Bullet')
                for run in p.runs:
                    run.font.size = Pt(9)

        # --- Notes ---
        if notes:
            doc.add_heading("Notes", level=2)
            p = doc.add_paragraph(str(notes).strip())
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Last updated footer
        if last_updated:
            p = doc.add_paragraph(f"Last updated: {last_updated}")
            p.paragraph_format.space_before = Pt(6)
            for run in p.runs:
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Page break between actors (except last)
        if actor_idx < len(actor_items) - 1:
            doc.add_page_break()

    return doc


# --- Main ---

def main():
    parser = argparse.ArgumentParser(description="Generate DOCX documents from research data.")
    parser.add_argument('--languages', action='store_true', help='Generate languages document only')
    parser.add_argument('--actors', action='store_true', help='Generate actors document only')
    args = parser.parse_args()

    # If neither flag, generate both
    do_languages = args.languages or (not args.languages and not args.actors)
    do_actors = args.actors or (not args.languages and not args.actors)

    # Load data
    print("Loading data...")
    languages = load_all_languages()
    actors = load_all_actors()
    focus_codes = load_focused_languages()

    # Filter to focused languages only
    focus_languages = {code: languages[code] for code in focus_codes if code in languages}

    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if do_languages:
        print(f"Generating languages document ({len(focus_languages)} languages)...")
        doc = generate_languages_doc(focus_languages, actors)
        out_path = OUTPUT_DIR / "WCA_NLP_Languages.docx"
        doc.save(str(out_path))
        print(f"  Saved: {out_path}")

    if do_actors:
        actor_count = len([k for k in actors if k != 'actor-template'])
        print(f"Generating actors document ({actor_count} actors)...")
        doc = generate_actors_doc(actors, focus_languages)
        out_path = OUTPUT_DIR / "WCA_NLP_Actors.docx"
        doc.save(str(out_path))
        print(f"  Saved: {out_path}")

    print("Done.")


if __name__ == '__main__':
    main()
